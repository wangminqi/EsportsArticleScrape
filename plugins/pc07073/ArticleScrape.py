import requests
import re
from docx import Document
from docx.shared import Inches
import docx.image.exceptions
import urllib


def collect_raw_content(url):
    response = requests.get(url)
    response.encoding = "utf-8"
    return response

def refine_content(response):
    pattern_1 = '<div class="lt_con">[\s\S]*?<div class="lt_con">'
    try:
	target_1 = re.search(pattern_1, response.text, flags=re.S).group()
    except AttributeError:
	return None
    target_1 = re.sub('<div class="lt_con">[\s\S]*?<h1>','',target_1)
    target_1 = re.sub('</h1>[\s\S]*?</script>','',target_1)
    target_1 = re.sub('<a href="[\s\S]*?">', '', target_1, count=0, flags=0)
    target_1 = re.sub('<span style[\s\S]*?>','',target_1)
    target_1 = re.sub('&nbsp;','',target_1)
    target_1 = re.sub('&lt;','',target_1)
    target_1 = re.sub('&gt;','',target_1)
    target_1 = re.sub('</strong>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<strong>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</span>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<br/>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<br />', '', target_1, count=0, flags=0)
    target_1 = re.sub('<p style="[\s\S]*?">','',target_1)
    target_1 = re.sub('<p>','',target_1)
    target_1 = re.sub('</p>','',target_1)
    target_1 = re.sub('</div>','',target_1)
    target_1 = re.sub('<div class="list">','',target_1)
    target_1 = re.sub('<!--[\s\S]*?-->','',target_1)
    target_1 = re.sub('</a>','',target_1)
    target_1 = re.sub('<table class="table2">[\s\S]*?<div class="lt_con">','',target_1)
    target_1 = re.sub('<div class="lt_con">','',target_1)
    target_1 = re.sub('<p align="center">','',target_1)
    return target_1

def document_append(document, content, url):
    document.add_paragraph(url)
    temp_content = content.split('<')
    for teamp1 in teamp_content:
        teamp_conten=teamp1.split('>')
        for i in teamp_conten:
            # 长度不为空
            if len(i) >= 1:
            # 确认非 照片链接
                if 'src="' in i:
                    # 如果是包含照片链接，下载到本地，写入word中
                    image_url = re.search('src="[^ ]*"', i, flags=re.S).group()[5:-1]
                    print(image_url)
                    if 'http' not in image_url:
                        image_url='http:'+image_url
                    if __name__ == '__main__':
                        path = '../../images/'
                    else:
                        path = './images/'
                    image_name = image_url.split('/')[-1]
                    try:
# 将 urllib.request.urlretrieve(image_url, path + image_name) 移动至try内
                        urllib.request.urlretrieve(image_url, path + image_name)
                        document.add_picture(path + image_name, width=Inches(5))
                        document.add_paragraph(image_url)
                        document.add_paragraph(image_name)
                    except docx.image.exceptions.UnrecognizedImageError:
                        document.add_paragraph('无效照片!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!' + image_name)
                    except docx.image.exceptions.UnexpectedEndOfFileError:
                        document.add_paragraph('图片有效但格式特殊，请手动插入 ' + image_name + '!!!!!!!!!!!!!!!!!!!')
# 增加OSError以判断是否爬取非图片资料
                    except OSError:
                        document.add_paragraph('文章中可能包含视频，请检查 ' + image_name + '!!!!!!!!!!!!!!!!!!!')
                else:
                    # 不包含图片链接，则直接写入word
                    document.add_paragraph(i)
            else:
                document.add_paragraph(i)
    document.add_paragraph(" ")
    document.add_paragraph(" ")
    document.add_paragraph(" ")
    return document

if __name__ == '__main__':
    document = Document()
    document.add_heading('Document Title', 0)

    # 读取并处理url
    with open('07073.txt', 'r') as f:
        a = f.readlines()
    b = []
    for i in a:
        b.append(i[:-1])

    content = []
    for url in b:
        # 正则处理
        response = collect_raw_content(url)
        target_1 = refine_content(response)
        print(target_1)
        content.append([target_1, url])
    for temp in content:
        document_append(document, temp[0], temp[1])

    # document.add_paragraph(target_1)
    # document.add_paragraph(url)

    document.save('07073.docx')



