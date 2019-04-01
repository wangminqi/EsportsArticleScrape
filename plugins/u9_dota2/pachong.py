import requests
import re
from docx import Document
from docx.shared import Inches
import docx.image.exceptions
import urllib
document = Document()
document.add_heading('Document Title', 0)
with open('uu9.txt', 'r') as f:
    a = f.readlines()
b = []
for i in a:
    b.append(i[:-1])

content = []
for url in b:
 
    # 正则处理
    response = requests.get(url)
    response.encoding = "gb2312"
    pattern_1 = '<div class="detail content">[\s\S]*?<p style="text-align:center; height:10px; margin-top:10px;">'
    target_1 = re.search(pattern_1, response.text, flags=re.S).group()
    target_1 = re.sub('<div style="[\s\S]*?">','',target_1)
    target_1=re.sub('<h3>[\s\S]*?<div class="textdetail" id="content222">','',target_1)
    target_1 = re.sub('<p style="box-sizing:[\s\S]*?">','', target_1, count=0, flags=0)
    target_1 = re.sub('<p style="margin-top:[\s\S]*?">','', target_1, count=0, flags=0)
    target_1 = re.sub('<span style="[\s\S]*?">', '', target_1, count=0, flags=0)
    target_1 = re.sub('<a href="[\s\S]*?">', '', target_1, count=0, flags=0)
    target_1 = re.sub('<p img-box="img-box[\s\S]*?">','', target_1, count=0, flags=0)
    target_1 = re.sub('<p>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</a>','',target_1, count=0, flags=0)
    target_1 = re.sub('</span>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<br/>','',target_1, count=0, flags=0)
    target_1 = re.sub(
        '<div class="article">[\s\S]*<h1>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</h1>[\s\S]*?<td>', '', target_1, count=0,
                      flags=0)
    target_1 = re.sub('<div class="journalism-code" style="color:red;">', '',
                      target_1,count=0, flags=0)
    target_1 = re.sub('&nbsp;','',target_1, count=0, flags=0)
    target_1 = re.sub('</td>[\s\S]*?<div class="article-tag clearfix">','',target_1)
    target_1 = re.sub('</div>','',target_1)
    target_1 = re.sub('<p style="text-indent:2em;">','',target_1)
    target_1 = re.sub('<p style="TEXT-INDENT: 2em">','',target_1)
    target_1 = re.sub('<strong>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</strong>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<p align="center">', '', target_1)
    target_1 = re.sub('</p>', '', target_1, count=0, flags=0)
  #  target_1 = re.sub('<a href="[^ ]*" target="_blank" >', '', target_1, count=0, flags=0)
    target_1 = re.sub('<a target="[\s\S]*?">', '', target_1, count=0, flags=0)
    target_1=re.sub('<p style="text-align[\s\S]*?">','',target_1)
    target_1=re.sub('<h1>','',target_1)
    target_1=re.sub('</h1>','',target_1)
    target_1=re.sub('<div class="detail content">','',target_1)
    content.append([target_1,url])
    print(url)

for temp in content:
    # 单篇文章分类
    document.add_paragraph(temp[1])
    temp_content = temp[0].split('\n')
    for i in temp_content:
        # 长度不为空
        if len(i) >= 1:
            # 确认非 照片链接
            if i[0] == '<':
                if 'src="' in i:
                    # 如果是包含照片链接，下载到本地，写入word中
                    image_url = re.search('src="[^ ]*"', i, flags=re.S).group()[5:-1]
                    path ='/home/yanziao/文档/工作/info/image/dota2/uuu9/'
                    image_name = image_url.split('/')[-1]

                    urllib.request.urlretrieve(image_url, path + image_name)
                    try:
                        document.add_picture(path + image_name, width=Inches(5))
                        document.add_paragraph(image_url)
                        document.add_paragraph(image_name)
                    except docx.image.exceptions.UnrecognizedImageError:
                        document.add_paragraph('无效照片!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!' + image_name)
                    except docx.image.exceptions.UnexpectedEndOfFileError:
                        document.add_paragraph('图片有效但格式特殊，请手动插入 ' + image_name + '!!!!!!!!!!!!!!!!!!!')
                else:
                    # 不包含图片链接，则直接写入word
                    document.add_paragraph(i)
            else:
                document.add_paragraph(i)
    document.add_paragraph(" ")
    document.add_paragraph(" ")
    document.add_paragraph(" ")

# document.add_paragraph(target_1)
# document.add_paragraph(url)


document.save('uu9.docx')
