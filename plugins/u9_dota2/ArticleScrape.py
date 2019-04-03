import requests
import re
from docx import Document
from docx.shared import Inches
import docx.image.exceptions
import urllib


def collect_raw_content(url):
    response = requests.get(url)
    response.encoding = "gb2312"
    return response


def refine_content(response):
    pattern_1 = '<div class="detail content">[\s\S]*?<p style="text-align:center; height:10px; margin-top:10px;">'
    try:
        target_1 = re.search(pattern, response.text, flags=re.S).group()
    except AttributeError:
        return None
    target_1 = re.sub('<div style="[\s\S]*?">','',target_1)
    target_1=re.sub('<h3>[\s\S]*?<div class="textdetail" id="content222">','',target_1)
    target_1 = re.sub('<p style="box-sizing:[\s\S]*?">','', target_1, count=0, flags=0)
    target_1 = re.sub('<p style="margin-top:[\s\S]*?">','', target_1, count=0, flags=0)
    target_1 = re.sub('<span style="[\s\S]*?">', '', target_1, count=0, flags=0)
    target_1 = re.sub('<a href="[\s\S]*?">', '', target_1, count=0, flags=0)
    target_1 = re.sub('<p img-box="img-box[\s\S]*?">','', target_1, count=0, flags=0)
    target_1 = re.sub('<p>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</a>','',target_1, count=0, flags=0)
    target_1 = re.sub('</span>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<br/>','',target_1, count=0, flags=0)
    target_1 = re.sub(
        '<div class="article">[\s\S]*<h1>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</h1>[\s\S]*?<td>', '', target_1, count=0,
                      flags=0)
    target_1 = re.sub('<div class="journalism-code" style="color:red;">', '',
                      target_1,count=0, flags=0)
    target_1 = re.sub('&nbsp;','',target_1, count=0, flags=0)
    target_1 = re.sub('</td>[\s\S]*?<div class="article-tag clearfix">','',target_1)
    target_1 = re.sub('</div>','',target_1)
    target_1 = re.sub('<p style="text-indent:2em;">','',target_1)
    target_1 = re.sub('<p style="TEXT-INDENT: 2em">','',target_1)
    target_1 = re.sub('<strong>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</strong>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<p align="center">', '', target_1)
    target_1 = re.sub('</p>', '', target_1, count=0, flags=0)
  #  target_1 = re.sub('<a href="[^ ]*" target="_blank" >', '', target_1, count=0, flags=0)
    target_1 = re.sub('<a target="[\s\S]*?">', '', target_1, count=0, flags=0)
    target_1=re.sub('<p style="text-align[\s\S]*?">','',target_1)
    target_1=re.sub('<h1>','',target_1)
    target_1=re.sub('</h1>','',target_1)
    target_1=re.sub('<div class="detail content">','',target_1)
    return target_1


def document_append(document, content, url):
    document.add_paragraph(url)
    temp_content = content.split('\n')
    for i in temp_content:
        # 长度不为空
        if len(i) >= 1:
            # 确认非 照片链接
            if i[0] == '<' or ' ':
                if 'src=' in i:
                    # 如果是包含照片链接，下载到本地，写入word中
                    image_url = re.search('src="[\s\S]*?"', i, flags=re.S).group()[5:-1]
                    print(image_url)
# 增加路径使用方式判断
                    if __name__ == '__main__':
                        path = '../../images/'
                    else:
                        path = './images/'
                    image_name = image_url.split('/')[-1]
                    try:
# 将 urllib.request.urlretrieve(image_url, path + image_name) 移动至try内
                        urllib.request.urlretrieve(image_url, path + image_name)
                        document.add_picture(path + image_name, width=Inches(5))
                        document.add_paragraph(image_url)
                        document.add_paragraph(image_name)
                    except docx.image.exceptions.UnrecognizedImageError:
                        document.add_paragraph('无效照片!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!' + image_name)
                    except docx.image.exceptions.UnexpectedEndOfFileError:
                        document.add_paragraph('图片有效但格式特殊，请手动插入 ' + image_name + '!!!!!!!!!!!!!!!!!!!')
# 增加OSError以判断是否爬取非图片资料
                    except OSError:
                        document.add_paragraph('文章中可能包含视频，请检查 ' + image_name + '!!!!!!!!!!!!!!!!!!!')
                else:
                    # 不包含图片链接，则直接写入word
                    document.add_paragraph(i)
            else:
                document.add_paragraph(i)
    document.add_paragraph(" ")
    document.add_paragraph(" ")
    document.add_paragraph(" ")
# 增加 return document
    return document


if __name__ == '__main__':
    document = Document()
    document.add_heading('Document Title', 0)

    # 读取并处理url
    with open('uu9.txt', 'r') as f:
        a = f.readlines()
    b = []
    for i in a:
        b.append(i[:-1])

    content = []
    for url in b:
        # 正则处理
        response = collect_raw_content(url)
        target_1 = refine_content(response)
        print(target_1)
        content.append([target_1, url])
    for temp in content:
        document_append(document, temp[0], temp[1])

    # document.add_paragraph(target_1)
    # document.add_paragraph(url)

    document.save('uuu9.docx')



