import requests
import re
from docx import Document
from docx.shared import Inches
import docx.image.exceptions
import urllib

document = Document()
document.add_heading('Document Title', 0)

# 读取并处理url
with open('csgo.txt', 'r') as f:
    a = f.readlines()
b = []
for i in a:
    b.append(i[:-1])

content = []
for url in b:
# 正则处理
    response = requests.get(url)
    response.encoding = "utf-8"
    pattern_1 = '<div class="new_header">[\s\S]*?<div class="page_cms">'
    target_1 = re.search(pattern_1, response.text, flags=re.S).group()
    target_1 = re.sub('<div class="new_header">[\s\S]*?<h2>','',target_1)
    target_1 = re.sub('</h2>[\s\S]*?<span style[\s\S]*?>','',target_1)
    target_1 = re.sub('<span style[\s\S]*?>','',target_1)
    target_1 = re.sub('&nbsp;','',target_1)
    target_1 = re.sub('&lt;','',target_1)
    target_1 = re.sub('&gt;','',target_1)
    target_1 = re.sub('</strong>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<strong>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</span>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<br/>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<br />', '', target_1, count=0, flags=0)
    target_1 = re.sub('<p style="[\s\S]*?">','',target_1)
    target_1 = re.sub('</p>','',target_1)
    target_1 = re.sub('</div>','',target_1)
    target_1 = re.sub('<div class="page_cms">','',target_1)
    target_1 = re.sub('<a href=[\s\S]*?>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<b>','',target_1)
    target_1 = re.sub('</b>','',target_1)
    target_1 = re.sub('<a style[\s\S]*?>','',target_1)
    target_1 = re.sub('</a>','',target_1)
    target_1 = re.sub('<p>','',target_1)
    target_1 = re.sub('<b style=[\s\S]*?>','',target_1)
    content.append([target_1, url])
for temp in content:
    document.add_paragraph(temp[1])
    teamp_content=temp[0].split('<')
    for teamp1 in teamp_content:
        teamp_conten=teamp1.split('>')
        for i in teamp_conten:
            print(i)
            # 长度不为空
            if len(i) >= 1:
                # 确认非 照片链接
                if 'src="' in i:
                    # 如果是包含照片链接，下载到本地，写入word中
                    image_url = re.search('src="[^ ]*"', i, flags=re.S).group()[5:-1]
                    path ='/home/yanziao/文档/工作/info/image/csgo/官网/'
                    image_name = image_url.split('/')[-1]
                    urllib.request.urlretrieve(image_url, path + image_name)
                    try:
                        document.add_picture(path + image_name, width=Inches(5))
                        document.add_paragraph(image_url)
                        document.add_paragraph(image_name)
                    except docx.image.exceptions.UnrecognizedImageError:
                        document.add_paragraph('无效照片!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!' + image_name)
                    except docx.image.exceptions.UnexpectedEndOfFileError:
                        document.add_paragraph('图片有效但格式特殊，请手动插入 ' + image_name + '!!!!!!!!!!!!!!!!!!!')
                else:
                    # 不包含图片链接，则直接写入word
                    document.add_paragraph(i)
            else:
                document.add_paragraph(i)
    document.add_paragraph(" ")
    document.add_paragraph(" ")
    document.add_paragraph(" ")

# document.add_paragraph(target_1)
# document.add_paragraph(url)


document.save('csgo官网.docx')
