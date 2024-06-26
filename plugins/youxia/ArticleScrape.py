import requests
import re
from docx import Document
from docx.shared import Inches
import docx.image.exceptions
import urllib


def collect_raw_content(url):
    response = requests.get(url)
    response.encoding = "utf-8"
    return response


def refine_content(response):
    pattern = '<div class="ns_t4">[\s\S]*?<div class="news_ding">'
    try:
        target = re.search(pattern, response.text, flags=re.S).group()
    except AttributeError:
        return None
    target = re.sub('<div class="ns_t4">[\s\S]*?<h1 class="newstit">', '', target)
    target = re.sub('</h1>[\s\S]*?<h2>', '', target)
    target = re.sub('</h2>[\s\S]*?</script>', '', target)
    target = re.sub('<p style="[\s\S]*?">', '', target)
    target = re.sub('<span style="[^\t\r\n]+">', '', target, count=0, flags=0)
    target = re.sub('<a href=[\s\S]*?>', '', target, count=0, flags=0)
    target = re.sub('<span class="n_show_g">[\s\S]*<div class="news_ding">', '', target, count=0, flags=0)
    target = re.sub('<p>', '', target)
    target = re.sub('</p>', '', target)
    target = re.sub('<strong>', '', target)
    target = re.sub('</strong>', '', target)
    target = re.sub('</span>', '', target, count=0, flags=0)
    target = re.sub('<div class="n_show" id="Content">', '', target)
    target = re.sub('<script src="[\s\S]*?</script>', '', target)
    target = re.sub('<font [\s\S]*?>', '', target)
    target = re.sub('<b>', '', target)
    target = re.sub('</b>', '', target)
    target = re.sub('</a>', '', target)
    target = re.sub('</font>', '', target)
    target = re.sub('&ldquo;', '', target)
    target = re.sub('&rdquo;', '', target)
    target = re.sub('&mdash;', '', target)
    target = re.sub('<div class="n_show_xg">[\s\S]*?<div class="news_ding">', '', target)
    target = re.sub('<script>[\s\S]*?<div class="news_ding">', '', target)
    return target


def document_append(document, content, url):
    document.add_paragraph(url)
    temp_content = content.split('\n')
    for i in temp_content:
        # 长度不为空
        if len(i) >= 1:
            # 确认非 照片链接
            if i[0] == '<' or ' ':
                if 'src=' in i:
                    # 如果是包含照片链接，下载到本地，写入word中
                    image_url = re.search('src="[\s\S]*?"', i, flags=re.S).group()[5:-1]
                    print(image_url)
# 增加路径使用方式判断
                    if __name__ == '__main__':
                        path = '../../images/'
                    else:
                        path = './images/'
                    image_name = image_url.split('/')[-1]
                    try:
# 将 urllib.request.urlretrieve(image_url, path + image_name) 移动至try内
                        urllib.request.urlretrieve(image_url, path + image_name)
                        document.add_picture(path + image_name, width=Inches(5))
                        document.add_paragraph(image_url)
                        document.add_paragraph(image_name)
                    except docx.image.exceptions.UnrecognizedImageError:
                        document.add_paragraph('无效照片!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!' + image_name)
                    except docx.image.exceptions.UnexpectedEndOfFileError:
                        document.add_paragraph('图片有效但格式特殊，请手动插入 ' + image_name + '!!!!!!!!!!!!!!!!!!!')
# 增加OSError以判断是否爬取非图片资料
                    except OSError:
                        document.add_paragraph('文章中可能包含视频，请检查 ' + image_name + '!!!!!!!!!!!!!!!!!!!')
                else:
                    # 不包含图片链接，则直接写入word
                    document.add_paragraph(i)
            else:
                document.add_paragraph(i)
    document.add_paragraph(" ")
    document.add_paragraph(" ")
    document.add_paragraph(" ")
# 增加 return document
    return document


if __name__ == '__main__':
    document = Document()
    document.add_heading('Document Title', 0)

    # 读取并处理url
    with open('堡垒之夜.txt', 'r') as f:
        a = f.readlines()
    b = []
    for i in a:
        b.append(i[:-1])

    content = []
    for url in b:
        # 正则处理
        response = collect_raw_content(url)
        target_1 = refine_content(response)
        print(target_1)
        content.append([target_1, url])
    for temp in content:
        document_append(document, temp[0], temp[1])

    # document.add_paragraph(target_1)
    # document.add_paragraph(url)

    document.save('堡垒之夜游侠.docx')



