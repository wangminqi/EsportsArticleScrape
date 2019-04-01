import requests
import re
from docx import Document
from docx.shared import Inches
import docx.image.exceptions
import urllib
import os
import time

from selenium import webdriver
from selenium.webdriver.chrome.options import Options


webdriver_path = '/home/yanziao/下载/chromedriver'
Chrome_path = '/opt/google/chrome/google-chrome'
options = Options()
options.binary_location = Chrome_path
options.add_argument('headless')
driver=webdriver.Chrome(executable_path=webdriver_path,options=options)
a = driver.find_elements_by_tag_name('p')
document = Document()
document.add_heading('Document Title', 0)
dir=os.getcwd()
print(dir)

# 读取并处理url
with open('loltest.txt', 'r') as f:
    a = f.readlines()
b = []
for i in a:
    b.append(i[:-1])
x=0
content = []
for url in b:
    driver.get(url)
    time.sleep(0.1)
    pattern_1 = '<div class="in-part">[\s\S]*?<div class="like-box">'
    target_1 = re.search(pattern_1, driver.page_source, flags=re.S).group()
    target_1 = re.sub('<strong>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</strong>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<p>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</p>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<a>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</a>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<span>', '', target_1, count=0, flags=0)
    target_1 = re.sub('</span>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<a href="[^ ]*" target="_blank" >', '', target_1, count=0, flags=0)
    target_1 = re.sub('<span style="[^\t\r\n]+">', '', target_1, count=0, flags=0)
    target_1 = re.sub('<div id="text">\s*<div class="ZQ-gap">\s*</div>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<div class="ZQ-gap">\s*</div>\s*</div>', '', target_1, count=0, flags=0)
    target_1 = re.sub('<div class="in-part">[\s\S]*<h1 class="art-tit">', '', target_1, count=0, flags=0)
    target_1 = re.sub('</h1>[\s\S]*<div class="article" id="article">','',target_1,count=0,
                      flags=0)
    target_1 = re.sub('</div>','',target_1)
    target_1 = re.sub('<div class="like-box">','',target_1)
    target_1 = re.sub('<p style="[\s\S]*?">','',target_1)
    target_1 = re.sub('\t', '', target_1, count=0, flags=0)
    target_1 = re.sub('\r', '', target_1, count=0, flags=0)
    target_1 = re.sub('\u3000', '', target_1, count=0, flags=0)
    target_1 = re.sub('  +', '', target_1, count=0, flags=0)
    content.append([target_1, url])
    x+=1
    print(target_1)
for temp in content:
    document.add_paragraph(temp[1])
    temp_content=temp[0].split('\n')
    for i in temp_content:
        # 长度不为空
        if len(i) >= 1:
            # 确认非照片链接
            if i[0] == '<':
                if 'src="' in i:
                    print(i)
                    # 如果是包含照片链接，下载到本地，写入word中
                    image_url = re.search('src="[\s\S]*?"', i, flags=re.S).group()[5:-1]
                    print(image_url)
                    path ='/home/yanziao/文档/工作/info/image/lol/lol官网/'
                    image_name = image_url.split('/')[-2]
                    if "http" in image_url:
                        urllib.request.urlretrieve(image_url, path + image_name)
                        try:
                            document.add_picture(path + image_name, width=Inches(5))
                            document.add_paragraph(image_url)
                            document.add_paragraph(image_name)
                        except docx.image.exceptions.UnrecognizedImageError:
                            document.add_paragraph('无效照片!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!' + image_name)
                        except docx.image.exceptions.UnexpectedEndOfFileError:
                            document.add_paragraph('图片有效但格式特殊，请手动插入 ' + image_name + '!!!!!!!!!!!!!!!!!!!')
                else:
                    # 不包含图片链接，则直接写入word
                    document.add_paragraph(i)
            else:
                document.add_paragraph(i)
    document.add_paragraph(" ")
    document.add_paragraph(" ")
    document.add_paragraph(" ")


document.save('lol官网.docx')

